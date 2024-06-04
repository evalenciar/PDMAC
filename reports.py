# -*- coding: utf-8 -*-
"""
Generate Reports
Created on 09/20/2022

The objective of this script is to generate reports based on the results.

@author: evalencia
"""

# Report modules
from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches

templates_path = 'templates/'

def All(results_path, int_review_path, report_dent_info, report_strain, report_RLA, graph_paths):
    # Template Settings
    report_temp = templates_path + 'Output Report All.docx'
    
    table_info = [#T    R   C   Height
                  (1,   4,  0,  1.5),
                  (1,   4,  1,  1.5),
                  (1,   6,  0,  1.75),
                  (1,   7,  0,  1.75),
                  (2,   7,  0,  2),
                  (2,   7,  1,  2),
                  (2,   9,  0,  2),
                  (2,   9,  1,  2),
                  (2,   11, 0,  2.75)]
    
    # height_strain = [1.5, 1.5, 1.75, 1.75]
    # height_RLA = [2, 2, 2, 2, 2.75]
    # graphs_strain = [int_review_path]
    # graphs_RLA = [int_review_path + '13_MPS_Dent.png',         # FEA Max. Principal Stress (psi)
    #               int_review_path + '14_Nodal_Path_OR.png',    # Max. Principal Stress and Dent Profile
    #               int_review_path + '12_ISO_ODMAX_OD.png']     # FEA Isometric View of Max. Principal Stress (psi)
    
    
    # Dent Information
    dent_ID, odometer, OD, WT, depth, length, width, orientation, pipe_grade, interaction = report_dent_info
    
    # Dent Strain Results
    eo_max, ei_max, strain_limit, strain_pass_fail = report_strain
    
    # Remaining Life Results
    SSI, CI, TTO5, MD49Class, SCF, MD49Life, SCFLife, MD49DesignLife, SCFDesignLife = report_RLA
    
    report_out  = results_path + 'Report_All_' + str(dent_ID) + '.docx'
    
    with MailMerge(report_temp) as document:
        # print(document.get_merge_fields())
        document.merge(
            # Dent Information
            DentNumber              = str(dent_ID),
            DentOdometer            = str(odometer),
            DentOD                  = str(round(OD,2)),
            DentWT                  = str(round(WT,3)),
            DentDepthI              = str(round(depth,3)),
            DentDepthP              = str(round(100*depth/(OD/2),3)),
            DentLength              = str(round(length,3)),
            DentWidth               = str(round(width,3)),
            DentOrientation         = str(orientation),
            PipeGrade               = str(pipe_grade),
            DentInteraction         = str(interaction),
            # Dent Strain Results
            StrainMaxOut            = str(round(eo_max*100,3)),
            StrainMaxIn             = str(round(ei_max*100,3)),
            StrainLimit             = str(round(strain_limit*100,3)),
            StrainPassFail          = str(strain_pass_fail),
            # Dent Remaining Life Results
            SSI                     = str(round(SSI,3)),
            CI                      = str(round(CI,3)),
            TTO5                    = str(TTO5),
            MD49Class               = str(MD49Class),
            SCF                     = str(round(SCF,3)),
            MD49Life                = str(round(MD49Life,3)),
            SCFLife                 = str(round(SCFLife,3)),
            MD49DesignLife          = str(round(MD49DesignLife,3)),
            SCFDesignLife           = str(round(SCFDesignLife,3)))
        document.write(report_out)
    
    # Insert images to the same report based on Table positions
    # Iterate through all of the graphs and save them to the document
    document = Document(report_out)
    tables = document.tables
    for i, val in enumerate(table_info):
        p = tables[val[0]].rows[val[1]].cells[val[2]].paragraphs[0]
        r = p.add_run()
        r.add_picture(graph_paths[i],height=Inches(val[3]))
    document.save(report_out)
    
    # # Dent Strain Results
    # # Dent Profile in Axial Direction
    # p = tables[1].rows[4].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_strain[0],height=Inches(height_strain[0]))
    # # Dent Profile in Circumferential Direction
    # p = tables[1].rows[4].cells[1].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_strain[1],height=Inches(height_strain[1]))
    # # Dent Strain Contour Map eo
    # p = tables[1].rows[6].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_strain[2],height=Inches(height_strain[2]))
    # # Dent Strain Contour Map ei
    # p = tables[1].rows[7].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_strain[3],height=Inches(height_strain[3]))
    
    # # Dent Remaining Life Results
    # # Pressure Cycle History
    # p = tables[2].rows[7].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_RLA[0],height=Inches(height_RLA[0]))
    # # Histogram
    # p = tables[2].rows[7].cells[1].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_RLA[1],height=Inches(height_RLA[1]))
    # # FEA Max. Principal Stress
    # p = tables[2].rows[9].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_RLA[2],height=Inches(height_RLA[2]))
    # # Max. Principal Stress and Dent Profile
    # p = tables[2].rows[9].cells[1].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_RLA[3],height=Inches(height_RLA[3]))
    # # FEA Isometric View of Max. Principal Stress
    # p = tables[2].rows[11].cells[0].paragraphs[0]
    # r = p.add_run()
    # r.add_picture(graphs_RLA[4],height=Inches(height_RLA[4]))
    # document.save(report_out)

def SCF(results_path, int_review_path, dent_ID, interaction, OD, WT, depth, length, width, MPs, SCF):
    
    report_temp = templates_path + 'Output Report SCF.docx'
    report_out  = results_path + 'Report_SCF_' + str(dent_ID) + '.docx'
    file_new_name = "ID" + str(dent_ID) + "_"
    
    with MailMerge(report_temp) as document:
        # print(document.get_merge_fields())
        document.merge(
            # Dent Information
            DentNumber              = str(dent_ID),
            Interaction             = str(interaction),
            DentOD                  = str(round(OD,2)),
            DentWT                  = str(round(WT,3)),
            DentDepthI              = str(round(depth,3)),
            DentDepthP              = str(round(100*depth/(OD/2),3)),
            DentLength              = str(round(length,3)),
            DentWidth               = str(round(width,3)),
            MaxPrincipalStress      = str(round(MPs,0)),
            SCF                     = str(round(SCF,3)),
            DentSeverity            = str("N/A"))
        document.write(report_out)
    # Insert images to the same report based on Table positions
    graph_stress_overall = int_review_path + file_new_name + '13_MPS_Dent.png'
    graph_stress_radius = int_review_path + file_new_name + '14_Nodal_Path_OR.png'
    graph_ISO = int_review_path + file_new_name + '12_ISO_ODMAX_OD.png'
    document = Document(report_out)
    tables = document.tables
    # Image 1: Circumferential Dent Profile
    p = tables[1].rows[2].cells[0].paragraphs[0]
    r = p.add_run()
    r.add_picture(graph_stress_overall, height=Inches(1.7))
    # Image 2: Longitudinal Dent Profile
    p = tables[1].rows[2].cells[1].paragraphs[0]
    r = p.add_run()
    r.add_picture(graph_stress_radius, height=Inches(2))
    # Image 3: 3D Dent Stress Field
    p = tables[1].rows[4].cells[0].paragraphs[0]
    r = p.add_run()
    r.add_picture(graph_ISO, height=Inches(2.7))
    document.save(report_out)
    